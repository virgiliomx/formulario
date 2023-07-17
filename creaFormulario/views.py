from io import StringIO
import os
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.templatetags.static import static
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl
from .models import Formulario

# Create your views here.


def formulario_list(request):
    formularios = Formulario.objects.all()
    print(formularios.count())

    paginator = Paginator(formularios, 10)
    page_number = request.GET.get('page', 1)
    try:
        formulario = paginator.page(page_number)
    except PageNotAnInteger:
        formulario = paginator.page(1)
    except EmptyPage:
        formulario = paginator.page(paginator.num_pages)

    return render(request, 'creaFormulario/index.html', {'formulario': formulario})


def cargaExcel(request):

    def boo(value):
        if value == "Sí":
            return True
        return False

    def dat(value):
        if value == "":
            return None
        return value

    if "GET" == request.method:
        return render(request, 'creaFormulario/cargaExcel.html', {})
    else:
        excel_file = request.FILES["excel_file"]

        # you may put validations here to check extension or file size

        wb = openpyxl.load_workbook(excel_file)

        hoja = wb.active
        # print(hoja)

        # reading a cell

        excel_data = []
        for i in range(2, hoja.max_row+1):
            # Hace split al campo de título, para no crear subcarpetas
            # titulo_e = (hoja.cell(i, 6).value).split("/", 1)
            # Valida valores boleanos para ser ingresados al modelo
            divulgacion_e = boo(hoja.cell(i, 13).value)
            invencion_e = boo(hoja.cell(i, 29).value)
            convenios_e = boo(hoja.cell(i, 31).value)
            mercado_e = boo(hoja.cell(i, 34).value)
            interes_e = boo(hoja.cell(i, 38).value)
            discovery_e = boo(hoja.cell(i, 40).value)
            compromiso_e = boo(hoja.cell(i, 43).value)
            contacto_e = boo(hoja.cell(i, 36).value)

            # Fechas
            fecha_1 = dat(hoja.cell(i, 15).value)
            fecha_2 = dat(hoja.cell(i, 16).value)
            fecha_3 = dat(hoja.cell(i, 17).value)

            formulario = Formulario(email_institucional=hoja.cell(i, 4).value,
                                    # titulo=titulo_e,
                                    titulo=hoja.cell(i, 7).value,
                                    dependencia=hoja.cell(i, 6).value,
                                    nombre=hoja.cell(i, 5).value,
                                    email=hoja.cell(i, 4).value,
                                    telefono=hoja.cell(i, 10).value,
                                    contacto=hoja.cell(i, 11).value,
                                    email_adicional=hoja.cell(i, 12).value,
                                    divulgacion=divulgacion_e,
                                    divulgacion_info=hoja.cell(i, 14).value,
                                    divulgacion_fecha_1=fecha_1,
                                    divulgacion_fecha_2=fecha_2,
                                    divulgacion_fecha_3=fecha_3,
                                    divulgacion_evidencia=hoja.cell(
                                        i, 18).value,
                                    keywords=hoja.cell(i, 19).value,
                                    area_aplicacion=hoja.cell(i, 20).value,
                                    tipo_invencion=hoja.cell(i, 21).value,
                                    descripcion=hoja.cell(i, 22).value,
                                    problematica=hoja.cell(i, 23).value,
                                    relevantes=hoja.cell(i, 24).value,
                                    ocurrencia=hoja.cell(i, 25).value,
                                    obvia=hoja.cell(i, 26).value,
                                    financiamiento=hoja.cell(i, 27).value,
                                    financiamiento_especifique=hoja.cell(
                                        i, 28).value,
                                    invencion_involucramiento=invencion_e,
                                    involucrameinto=hoja.cell(i, 30).value,
                                    convenios=convenios_e,
                                    convenios_tipo=hoja.cell(i, 32).value,
                                    necesidad=hoja.cell(i, 33).value,
                                    mercado=mercado_e,
                                    mercado_especifique=hoja.cell(i, 35).value,
                                    contacto_empresa=contacto_e,
                                    contacto_especifique=hoja.cell(
                                        i, 37).value,
                                    interes=interes_e,
                                    interes_especifique=hoja.cell(i, 39).value,
                                    discovery=discovery_e,
                                    discovery_especifique=hoja.cell(
                                        i, 41).value,
                                    informacion_tecnica=hoja.cell(i, 42).value,
                                    compromiso=compromiso_e)
            excel_data.append(formulario)
            print(formulario.divulgacion_fecha_1)
            formulario.save()
    return render(request, 'creaFormulario/cargaExcel.html', {"excel_data": excel_data})


def generaFormulario(request, formulario_id):

    formulario = Formulario.objects.get(pk=formulario_id)
    docx_title = "Formulario de Invención.docx"
    document = Document()
    ################################
    #           Header             #
    ################################
    header = document.sections[0].header
    pheader = header.paragraphs[0]
    logo_run = pheader.add_run()
    logo_run.add_picture(os.path.join(
        "static", "img/header.png"), height=Inches(1))
    logo_run.add_break()
    logo_run.alignment = 3

    footer = document.sections[0].footer
    pfooter = footer.paragraphs[0]
    footer_run = pfooter.add_run()
    footer_run.add_picture(os.path.join(
        "static", "img/footer.png"), height=Inches(1.5))
    footer_run.alignment = 3

    ################################
    #           Portada            #
    ################################
    p_tt = document.add_paragraph()
    run_tt = p_tt.add_run("COORDINACIÓN DE TRANSFERENCIA DE TECNOLOGÍA")
    p_pi = document.add_paragraph()
    run_pi = p_pi.add_run("COORDINACIÓN DE PROPIEDAD INDUSTRIAL")
    p_formulario = document.add_paragraph()
    run_formulario = p_formulario.add_run(
        "FORMULARIO DE DECLARACIÓN DE LA INVENCIÓN")
    run_formulario.bold = True
    p_tt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_pi.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_formulario.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_formulario.font.size = Pt(45)

    run_formulario.add_break()
    run_formulario.add_break()

    p_titulo = document.add_paragraph()
    run_titulo = p_titulo.add_run(
        f'Título del proyecto: {formulario.titulo}')
    run_titulo.add_break()
    run_nombre = p_titulo.add_run(
        f'Nombre: {formulario.nombre}, Tel: {formulario.telefono}, Correo: {formulario.email},')
    run_titulo.bold = True
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()

    p_divulgacion = document.add_paragraph()
    run_divulgacion = p_divulgacion.add_run("¿Ha realizado divulgación?")
    run_divulgacion.add_break()
    if not formulario.divulgacion:
        run_respuesta_divulgacion = p_divulgacion.add_run(
            f'{formulario.divulgacion}')
    else:
        run_respuesta_divulgacion = p_divulgacion.add_run(
            f'{formulario.divulgacion}\n{formulario.divulgacion_info} en fecha {formulario.divulgacion_fecha_1}')
        if formulario.divulgacion_fecha_2 is not None:
            run_fecha_divulgacion_2 = p_divulgacion.add_run(
                f', {formulario.divulgacion_fecha_2}')
        if formulario.divulgacion_fecha_3 is not None:
            run_fecha_divulgacion_3 = p_divulgacion.add_run(
                f', {formulario.divulgacion_fecha_3}')
        run_divulgacion_evidencia = p_divulgacion.add_run(
            f', evidencia {formulario.divulgacion_evidencia}.')
    run_divulgacion.bold = True

    p_keywords = document.add_paragraph()
    run_keywords = p_keywords.add_run(
        "Proporcione una lista de palabras claves técnicas en español e inglés que describan su tecnología:")
    run_keywords.add_break()
    run_respuesta_keywords = p_keywords.add_run(f'{formulario.keywords}')
    run_keywords.bold = True

    p_area_aplicacion = document.add_paragraph()
    run_area_aplicacion = p_area_aplicacion.add_run(
        "Mencione el área de aplicación de su invención:")
    run_area_aplicacion.add_break()
    run_respuesta_area_aplicacion = p_area_aplicacion.add_run(
        f'{formulario.area_aplicacion}')
    run_area_aplicacion.bold = True

    p_tipo = document.add_paragraph()
    run_tipo = p_tipo.add_run("Su invención es un:")
    run_tipo.add_break()
    run_respuesta_tipo = p_tipo.add_run(f'{formulario.tipo_invencion}')
    run_tipo.bold = True

    p_descripcion = document.add_paragraph()
    run_descripcion = p_descripcion.add_run(
        "Descripción detallada de su invención:")
    run_descripcion.add_break()
    run_despuesta_descripcion = p_descripcion.add_run(
        f'{formulario.descripcion}')
    run_descripcion.bold = True

    p_problematica = document.add_paragraph()
    run_problematica = p_problematica.add_run(
        "Mencione cuál es el problema técnico que resuelve su invención:")
    run_problematica.add_break()
    run_respuesta_problematica = p_problematica.add_run(
        f'{formulario.problematica}')
    run_problematica.bold = True

    p_relevantes = document.add_paragraph()
    run_relevantes = p_relevantes.add_run(
        "Mencione los artículos y/o patentes más relevantes:")
    run_relevantes.add_break()
    run_respuesta_relevantes = p_relevantes.add_run(
        f'{formulario.relevantes}')
    run_relevantes.bold = True

    p_ocurrencia = document.add_paragraph()
    run_ocurrencia = p_ocurrencia.add_run(
        "¿Cómo se le surgió la invención?")
    run_ocurrencia.add_break()
    run_respuesta_ocurrencia = p_ocurrencia.add_run(
        f'{formulario.ocurrencia}')
    run_ocurrencia.bold = True

    p_obvia = document.add_paragraph()
    run_obvia = p_obvia.add_run(
        "¿La invención podría resultar obvia para otra persona con conocimientos medios del área?, Especifique:")
    run_obvia.add_break()
    run_respuesta_obvia = p_obvia.add_run(f'{formulario.obvia}')
    run_obvia.bold = True

    p_financiamiento = document.add_paragraph()
    run_financiamiento = p_financiamiento.add_run(
        "Este proyecto fue financiado por:")
    run_financiamiento.add_break()
    run_respuesta_financiamiento = p_financiamiento.add_run(
        f'{formulario.financiamiento}')
    run_financiamiento.bold = True

    p_financiamiento_departamento = document.add_paragraph()
    run_financiamiento_departamento = p_financiamiento_departamento.add_run(
        "Especifique cual departamento:")
    run_financiamiento_departamento.add_break()
    run_respuesta_financiamiento_departamento = p_financiamiento_departamento.add_run(
        f'{formulario.financiamiento_especifique}')
    run_financiamiento_departamento.bold = True

    p_involucramiento = document.add_paragraph()
    run_involucramiento = p_involucramiento.add_run(
        "¿Existe alguna otra institución, organización o empresa involucrada en el desarrollo de la invención?, Especifique:")
    run_involucramiento.add_break()
    run_respuesta_involucramiento = p_involucramiento.add_run(
        f'{formulario.invencion_involucramiento}')
    if formulario.invencion_involucramiento:
        run_involucramiento_inv = p_involucramiento.add_run(
            f', {formulario.involucramiento}')
    run_involucramiento.bold = True

    p_convenios = document.add_paragraph()
    run_convenios = p_convenios.add_run(
        "¿Existen convenios o acuerdos con terceros involucrados?, Especifique:")
    run_convenios.add_break()
    run_respuesta_convenios = p_convenios.add_run(
        f'{formulario.convenios}')
    run_convenios.bold = True

    p_convenios_tipos = document.add_paragraph()
    run_convenios_tipos = p_convenios_tipos.add_run(
        "Especifique el tipo de convenio o acuerdo con terceros involucrados:")
    run_convenios_tipos.add_break()
    run_respuesta_convenios_tipos = p_convenios_tipos.add_run(
        f'{formulario.convenios_tipo}')
    run_convenios_tipos.bold = True

    p_necesidad = document.add_paragraph()
    run_necesidad = p_necesidad.add_run("La invención surgió a partir de:")
    run_necesidad.add_break()
    run_respuesta_necesidad = p_necesidad.add_run(
        f'{formulario.necesidad}')
    run_necesidad.bold = True

    p_mercado = document.add_paragraph()
    run_mercado = p_mercado.add_run(
        "Ha identificado el mercado de la invención:")
    run_mercado.add_break()
    run_respuesta_mercado = p_mercado.add_run(f'{formulario.mercado}')
    run_mercado.bold = True

    if formulario.mercado:
        p_mercado_especifique = document.add_paragraph()
        run_mercado_especifique = p_mercado_especifique.add_run(
            "Especifique:")
        run_mercado_especifique.add_break()
        run_respuesta_mercado_especifique = p_mercado_especifique.add_run(
            f'{formulario.mercado_especifique}')
        run_mercado_especifique.bold = True

    p_contacto = document.add_paragraph()
    run_contacto = p_contacto.add_run(
        "¿Se ha contactado con alguna empresa/socio/inversionista para su posible explotación?")
    run_contacto.add_break()
    run_respuesta_contacto = p_contacto.add_run(f'{formulario.contacto}')
    run_contacto.bold = True

    if formulario.contacto:
        p_contacto_especifique = document.add_paragraph()
        run_contacto_especifique = p_contacto_especifique.add_run(
            "Especifique:")
        run_contacto_especifique.add_break()
        run_respuesta_contacto_especifique = p_contacto_especifique.add_run(
            f'{formulario.contacto_especifique}')
        run_contacto_especifique.bold = True

    p_interesada = document.add_paragraph()
    run_interesada = p_interesada.add_run(
        "¿Conoce alguna empresa que pudiera estar interesada?")
    run_interesada.add_break()
    run_respuesta_interesada = p_interesada.add_run(
        f'{formulario.interes}')
    run_interesada.bold = True

    if formulario.interes:
        p_interesada_especifique = document.add_paragraph()
        run_interesada_especifique = p_interesada_especifique.add_run(
            "Especifique:")
        run_interesada_especifique.add_break()
        run_respuesta_interesada_especifique = p_interesada_especifique.add_run(
            f'{formulario.interes_especifique}')
        run_interesada_especifique.bold = True

    p_customer = document.add_paragraph()
    run_customer = p_customer.add_run(
        "¿Ha participado en algún proceso de customer discovery, validación de mercado, pre-incubación o programa de emprendimiento con esta invención?")
    run_customer.add_break()
    run_respuesta_customer = p_customer.add_run(f'{formulario.discovery}')
    run_customer.bold = True

    if formulario.discovery:
        p_ocurrencia = document.add_paragraph()
        run_ocurrencia = p_ocurrencia.add_run("Especifique:")
        run_ocurrencia.add_break()
        run_respuesta_ocurrencia = p_ocurrencia.add_run(
            f'{formulario.discovery_especifique}')
        run_ocurrencia.bold = True

    p_informacion_tecnica = document.add_paragraph()
    run_informacion_tecnica = p_informacion_tecnica.add_run(
        "Información técnica:")
    run_informacion_tecnica.add_break()
    run_respuesta_informacion_tecnica = p_informacion_tecnica.add_run(
        f'{formulario.informacion_tecnica}')
    run_informacion_tecnica.bold = True

    p_compromiso = document.add_paragraph()
    run_compromiso = p_compromiso.add_run(
        "Me comprometo a dar seguimiento a las acciones que me sean requeridas por el Centro de Incubación de Empresas y Transferencia de Tecnología con respecto a la protección y comercialización de la invención:")
    run_compromiso.add_break()
    run_respuesta_compromiso = p_compromiso.add_run(
        f'{formulario.compromiso}')
    run_compromiso.bold = True

    ################################
    # Genera documento para guardar#
    ################################
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    document.save(response)
    return response
